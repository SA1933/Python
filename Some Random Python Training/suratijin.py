from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Membuat dokumen baru
doc = Document()

# Mengatur margin
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Menambahkan header untuk nomor surat
header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.text = "Nomor: [Nomor Surat]"
header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Menambahkan judul dengan format bold
title = doc.add_paragraph('SURAT PERSETUJUAN ORANG TUA')
title_format = title.runs[0]
title_format.font.name = 'Times New Roman'
title_format.bold = True
title_format.font.size = Pt(14)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Menambahkan paragraf pembuka
doc.add_paragraph('\nYang bertanda tangan di bawah ini:')

# Menambahkan informasi orang tua/wali
parent_info = doc.add_paragraph()
parent_info.add_run('Nama: ').bold = True
parent_info.add_run('[Nama Lengkap Orang Tua/Wali]')

parent_info.add_run('\nTempat, Tanggal Lahir: ').bold = True
parent_info.add_run('[Tempat, Tanggal Lahir Orang Tua/Wali]')

parent_info.add_run('\nAlamat: ').bold = True
parent_info.add_run('[Alamat Lengkap Orang Tua/Wali]')

parent_info.add_run('\nNo. Telepon: ').bold = True
parent_info.add_run('[Nomor Telepon Orang Tua/Wali]')

# Menambahkan persetujuan
doc.add_paragraph('\nDengan ini memberikan persetujuan kepada anak kami:')

child_info = doc.add_paragraph()
child_info.add_run('Nama: ').bold = True
child_info.add_run('[Nama Lengkap Anda]')

child_info.add_run('\nTempat, Tanggal Lahir: ').bold = True
child_info.add_run('[Tempat, Tanggal Lahir Anda]')

child_info.add_run('\nAlamat: ').bold = True
child_info.add_run('[Alamat Lengkap Anda]')

child_info.add_run('\nNo. Telepon: ').bold = True
child_info.add_run('[Nomor Telepon Anda]')

# Menambahkan isi persetujuan
doc.add_paragraph('\nUntuk mengikuti program magang di Jepang yang diselenggarakan oleh [Nama Lembaga/Penyelenggara] selama [Durasi Magang, contoh: 6 bulan], mulai dari tanggal [Tanggal Mulai] sampai dengan tanggal [Tanggal Selesai].')

doc.add_paragraph('Kami sepenuhnya mendukung anak kami dalam mengikuti program ini dengan harapan dapat menambah pengalaman kerja, memperluas wawasan, serta meningkatkan kemampuan profesionalnya di bidang [Bidang Pekerjaan Anda]. Kami juga memahami bahwa program ini memerlukan tanggung jawab yang besar, dan kami telah memberikan pengarahan kepada anak kami untuk menjaga nama baik keluarga serta mengikuti semua peraturan yang berlaku selama program magang berlangsung.')

# Menambahkan penutup
doc.add_paragraph('\nDemikian surat persetujuan ini kami buat dengan penuh kesadaran dan tanpa ada paksaan dari pihak manapun. Atas perhatian dan kerja sama yang diberikan, kami ucapkan terima kasih.')

# Menambahkan tanda tangan
signature_paragraph = doc.add_paragraph('\nHormat kami,')
signature_paragraph.add_run('\n\n\n\n[TTD]').font.size = Pt(12)
signature_paragraph.add_run('\n[Nama Lengkap Orang Tua/Wali]').bold = True

# Mengatur font umum untuk seluruh dokumen
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Menyimpan dokumen
file_path = r"D:\kuliah\folder kuliah_AriqRasyaEkaMaulana\LPK"
doc.save(file_path)

file_path
